#!/usr/bin/env python3
"""Convert a marketing-copy export to .docx for Google Docs import.

Google Docs imports .docx with headings, bold runs, tables, and lists intact,
which markdown paste does not preserve. Requires python-docx.

Usage:
    python3 scripts/md-to-docx.py                      # default exports
    python3 scripts/md-to-docx.py path/to/file.md ...  # explicit files
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    sys.exit("python-docx is not installed. Run: python3 -m pip install --user python-docx")

ROOT = Path(__file__).resolve().parent.parent
COPY_DIR = ROOT / "docs" / "marketing-copy"

DEFAULT_EXPORTS = [
    COPY_DIR / "TASTEFIELD_HOMEPAGE_GOOGLE_DOCS.md",
    COPY_DIR / "TASTEFIELD_SITE_COPY_GOOGLE_DOCS.md",
]

# Inline markdown: **bold**, *italic*, `code`.
INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)")


def add_runs(paragraph, text: str) -> None:
    """Write text into a paragraph, converting inline markdown to real runs."""
    for token in INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        elif token.startswith("*") and token.endswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        else:
            paragraph.add_run(token)


def split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for column, value in enumerate(row):
            cell = cells[column]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_runs(paragraph, value)
            if index == 0:
                for run in paragraph.runs:
                    run.bold = True


def convert(md_path: Path) -> Path:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Helvetica"
    normal.font.size = Pt(11)

    pending_table: list[list[str]] = []

    def flush_table() -> None:
        if pending_table:
            add_table(doc, list(pending_table))
            doc.add_paragraph()
            pending_table.clear()

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            if not is_separator(stripped):
                pending_table.append(split_row(stripped))
            continue
        flush_table()

        if not stripped:
            continue

        if stripped == "---":
            doc.add_paragraph()
            continue

        heading = re.match(r"(#{1,6})\s+(.*)", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            paragraph = doc.add_heading(level=level)
            add_runs(paragraph, heading.group(2))
            continue

        numbered = re.match(r"\d+\.\s+(.*)", stripped)
        if numbered:
            add_runs(doc.add_paragraph(style="List Number"), numbered.group(1))
            continue

        bullet = re.match(r"[-*]\s+(.*)", stripped)
        if bullet:
            add_runs(doc.add_paragraph(style="List Bullet"), bullet.group(1))
            continue

        add_runs(doc.add_paragraph(), stripped)

    flush_table()

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


def main() -> None:
    targets = [Path(arg).resolve() for arg in sys.argv[1:]] or DEFAULT_EXPORTS
    for target in targets:
        if not target.exists():
            print(f"skip  {target} (not found)")
            continue
        out = convert(target)
        label = out.relative_to(ROOT) if out.is_relative_to(ROOT) else out
        print(f"wrote {label}  ({out.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
